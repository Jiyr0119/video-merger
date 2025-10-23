import os
import re
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


# === 配置区 ===
SOURCE_DIR = r"./docs"     # 输入目录
TARGET_DIR = r"./output"   # 输出目录
FONT_NAME = "宋体"          # 默认字体
FONT_SIZE = 12              # 统一字号


# 样式关键字智能识别规则
STYLE_KEYWORDS = {
    "heading_1": ["标题1", "一级标题", "主标题", "chapter", "heading1", "xx标题1"],
    "heading_2": ["标题2", "二级标题", "子标题", "heading2", "xx标题2"],
    "heading_3": ["标题3", "三级标题", "heading3"],
    "normal": ["正文", "正文样式", "body", "paragraph", "normaltext"],
}

AUTO_MAP = {
    "heading_1": "Heading 1",
    "heading_2": "Heading 2",
    "heading_3": "Heading 3",
    "normal": "Normal",
}


# === 功能函数 ===

def guess_new_style_by_name(style_name: str):
    """根据样式名猜测目标样式"""
    s = style_name.lower()
    for key, kws in STYLE_KEYWORDS.items():
        for kw in kws:
            if kw.lower() in s:
                return AUTO_MAP[key]
    return None


def guess_by_formatting(para):
    """
    根据段落格式判断标题层级
    - 字号 + 加粗比例判断
    """
    if not para.runs:
        return None

    sizes = []
    bold_count = 0
    for run in para.runs:
        if run.font.size:
            sizes.append(run.font.size.pt)
        if run.bold:
            bold_count += 1

    avg_size = sum(sizes) / len(sizes) if sizes else 0
    bold_ratio = bold_count / len(para.runs)

    if avg_size >= 16 and bold_ratio > 0.5:
        return "Heading 1"
    elif avg_size >= 13 and bold_ratio > 0.3:
        return "Heading 2"
    elif avg_size >= 11 and bold_ratio > 0.3:
        return "Heading 3"
    else:
        return "Normal"


def guess_by_numbering(text: str):
    """
    根据编号模式判断层级
    示例:
      "1 "      -> Heading 1
      "1.2 "    -> Heading 2
      "1.2.3 "  -> Heading 3
    """
    if not text:
        return None
    match = re.match(r"^(\d+(\.\d+){0,5})[\s、.:-]", text.strip())
    if match:
        num_str = match.group(1)
        level = num_str.count(".") + 1
        if level == 1:
            return "Heading 1"
        elif level == 2:
            return "Heading 2"
        elif level >= 3:
            return "Heading 3"
    return None


def set_font(run, font_name: str, font_size: int):
    """统一字体"""
    try:
        run.font.name = font_name
        r = run._element.rPr.rFonts
        r.set(qn('w:eastAsia'), font_name)
        r.set(qn('w:ascii'), font_name)
        r.set(qn('w:hAnsi'), font_name)
        run.font.size = Pt(font_size)
    except Exception:
        pass


def process_docx(filepath: str, target_dir: str):
    print(f"📝 正在处理: {filepath}")
    doc = Document(filepath)
    modified = False

    for para in doc.paragraphs:
        old_style = para.style.name
        text = para.text.strip()

        # Step 1: 样式名匹配
        new_style = guess_new_style_by_name(old_style)

        # Step 2: 编号模式匹配
        if not new_style:
            new_style = guess_by_numbering(text)

        # Step 3: 格式特征匹配
        if not new_style:
            new_style = guess_by_formatting(para)

        # Step 4: 应用样式
        if new_style and new_style != old_style:
            try:
                para.style = new_style
                modified = True
                print(f"  🟢 {text[:20]}... → {new_style}")
            except Exception as e:
                print(f"  ⚠️ 样式替换失败: {old_style} → {new_style} ({e})")

        # Step 5: 统一字体
        for run in para.runs:
            set_font(run, FONT_NAME, FONT_SIZE)

    # 表格内容同样处理
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    old_style = para.style.name
                    new_style = (
                        guess_new_style_by_name(old_style)
                        or guess_by_numbering(text)
                        or guess_by_formatting(para)
                    )
                    if new_style and new_style != old_style:
                        try:
                            para.style = new_style
                            modified = True
                        except Exception:
                            pass
                    for run in para.runs:
                        set_font(run, FONT_NAME, FONT_SIZE)

    # 保存新文件
    os.makedirs(target_dir, exist_ok=True)
    target_path = os.path.join(target_dir, os.path.basename(filepath))
    doc.save(target_path)
    if modified:
        print(f"✅ 已保存: {target_path}\n")
    else:
        print(f"ℹ️ 无修改: {os.path.basename(filepath)}\n")


def batch_process(source_dir: str, target_dir: str):
    for root, _, files in os.walk(source_dir):
        for file in files:
            if file.lower().endswith(".docx"):
                process_docx(os.path.join(root, file), target_dir)


if __name__ == "__main__":
    batch_process(SOURCE_DIR, TARGET_DIR)
    print("🎉 全部文件智能规范化完成！")
